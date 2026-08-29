"""Tests for aida.documents.tools.default_document_tools — the agent-facing
write_markdown_report/write_docx_report tools, including resolving
image_artifact_ids against a shared ArtifactStore's recorded metadata."""

from __future__ import annotations

from pathlib import Path

import pytest

from aida.artifacts.base import FileArtifact, ImageArtifact
from aida.artifacts.store import ArtifactStore
from aida.documents.tools import default_document_tools
from aida.workspace.safety import ConfirmationRequest, SafetyGuard
from tests.mock_mcp_server import TINY_PNG_BYTES


async def _approve(_request: ConfirmationRequest) -> bool:
    return True


def _guard(root: Path) -> SafetyGuard:
    return SafetyGuard(allowed_roots=[root], mode="relaxed", confirm_callback=_approve)


async def _call(tools, name: str, **arguments):
    return await tools[name].func(arguments)


@pytest.mark.asyncio
async def test_write_markdown_report_plain(tmp_path: Path):
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    tools = default_document_tools(_guard(tmp_path), store)
    result = await _call(
        tools, "write_markdown_report", path=str(tmp_path / "report.md"), title="My Report", body="Findings."
    )
    assert not result.is_error
    assert isinstance(result.artifacts[0], FileArtifact)
    text = Path(result.artifacts[0].path).read_text(encoding="utf-8")
    assert "# My Report" in text
    assert "Findings." in text


@pytest.mark.asyncio
async def test_write_markdown_report_with_known_image_artifact_id(tmp_path: Path):
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="plot.png"))
    tools = default_document_tools(_guard(tmp_path), store)

    result = await _call(
        tools,
        "write_markdown_report",
        path=str(tmp_path / "report.md"),
        title="With Plot",
        body="See below.",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    text = Path(result.artifacts[0].path).read_text(encoding="utf-8")
    assert "figures/" in text
    assert (tmp_path / "figures").exists()


@pytest.mark.asyncio
async def test_write_markdown_report_unknown_image_id_skipped_not_failed(tmp_path: Path):
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    tools = default_document_tools(_guard(tmp_path), store)
    result = await _call(
        tools,
        "write_markdown_report",
        path=str(tmp_path / "report.md"),
        title="Stale Ref",
        body="Body.",
        image_artifact_ids=["does-not-exist"],
    )
    assert not result.is_error
    text = Path(result.artifacts[0].path).read_text(encoding="utf-8")
    assert "Body." in text


@pytest.mark.asyncio
async def test_write_markdown_report_outside_allowed_folders_denied(tmp_path: Path):
    from aida.workspace.safety import deny_all

    allowed = tmp_path / "allowed"
    allowed.mkdir()
    outside = tmp_path / "outside.md"
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    guard = SafetyGuard(allowed_roots=[allowed], confirm_callback=deny_all)
    tools = default_document_tools(guard, store)
    result = await _call(tools, "write_markdown_report", path=str(outside), title="Nope")
    assert result.is_error
    assert not outside.exists()


@pytest.mark.asyncio
async def test_write_markdown_report_custom_sidecar_dirname(tmp_path: Path):
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="a.png"))
    tools = default_document_tools(_guard(tmp_path), store, sidecar_dirname="images")
    result = await _call(
        tools,
        "write_markdown_report",
        path=str(tmp_path / "report.md"),
        title="T",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    assert (tmp_path / "images").exists()
    assert not (tmp_path / "figures").exists()


@pytest.mark.asyncio
async def test_write_markdown_report_image_placeholder_places_it_inline(tmp_path: Path):
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="plot.png"))
    tools = default_document_tools(_guard(tmp_path), store)

    result = await _call(
        tools,
        "write_markdown_report",
        path=str(tmp_path / "report.md"),
        title="With Plot",
        body=f"Before.\n\n{{{{image:{image.id}}}}}\n\nAfter.",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    text = Path(result.artifacts[0].path).read_text(encoding="utf-8")
    assert text.index("Before.") < text.index("figures/") < text.index("After.")


# --- write_docx_report -------------------------------------------------------


@pytest.mark.asyncio
async def test_write_docx_report_plain(tmp_path: Path):
    pytest.importorskip("docx")
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    tools = default_document_tools(_guard(tmp_path), store)
    result = await _call(
        tools, "write_docx_report", path=str(tmp_path / "report.docx"), title="DOCX Title", body="Body text."
    )
    assert not result.is_error
    assert Path(result.artifacts[0].path).exists()

    import docx

    document = docx.Document(result.artifacts[0].path)
    texts = [p.text for p in document.paragraphs]
    assert "DOCX Title" in texts
    assert "Body text." in texts


@pytest.mark.asyncio
async def test_write_docx_report_with_image(tmp_path: Path):
    docx = pytest.importorskip("docx")
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="plot.png"))
    tools = default_document_tools(_guard(tmp_path), store)
    result = await _call(
        tools,
        "write_docx_report",
        path=str(tmp_path / "report.docx"),
        title="With Image",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    document = docx.Document(result.artifacts[0].path)
    assert len(document.inline_shapes) == 1

@pytest.mark.asyncio
async def test_write_docx_report_image_placeholder_places_it_between_paragraphs(tmp_path: Path):
    docx = pytest.importorskip("docx")
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="plot.png"))
    tools = default_document_tools(_guard(tmp_path), store)

    result = await _call(
        tools,
        "write_docx_report",
        path=str(tmp_path / "report.docx"),
        title="With Plot",
        body=f"Before.\n\n{{{{image:{image.id}}}}}\n\nAfter.",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    document = docx.Document(result.artifacts[0].path)
    assert len(document.inline_shapes) == 1

    texts = [p.text for p in document.paragraphs]
    before_idx = next(i for i, t in enumerate(texts) if "Before." in t)
    after_idx = next(i for i, t in enumerate(texts) if "After." in t)
    picture_idx = next(i for i, p in enumerate(document.paragraphs) if "graphicData" in p._p.xml)
    assert before_idx < picture_idx < after_idx


@pytest.mark.asyncio
async def test_write_docx_report_unreferenced_image_still_appended_after_body(tmp_path: Path):
    docx = pytest.importorskip("docx")
    store = ArtifactStore(base_dir=tmp_path / "artifacts")
    image = store.save_image(ImageArtifact(data=TINY_PNG_BYTES, mime_type="image/png", filename="plot.png"))
    tools = default_document_tools(_guard(tmp_path), store)

    result = await _call(
        tools,
        "write_docx_report",
        path=str(tmp_path / "report.docx"),
        title="No Placeholder",
        body="Just some text, no placeholder here.",
        image_artifact_ids=[image.id],
    )
    assert not result.is_error
    document = docx.Document(result.artifacts[0].path)
    assert len(document.inline_shapes) == 1
