"""Basic DOCX writer (PLAN.md Phase 6): "DOCX writer (basic: headings,
paragraphs, images, tables) for Office needs" — via ``python-docx`` (the
``docs`` extra). Not the default writer (``md_obsidian.py`` is); this is
specifically for Word/Office-centric output when Markdown isn't the right
fit for the user's workflow.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from aida.artifacts.base import ImageArtifact
from aida.workspace.safety import unique_destination


@dataclass
class DocxSection:
    """One piece of a document body, in the order they should appear."""

    kind: str  # "heading" | "paragraph" | "image" | "table"
    text: str = ""
    level: int = 1  # heading level (1-9); ignored for other kinds
    image: ImageArtifact | None = None
    columns: list[str] = field(default_factory=list)
    rows: list[list[Any]] = field(default_factory=list)


def write_docx_document(*, target_dir: Path, filename_stem: str, title: str, sections: list[DocxSection]) -> Path:
    """Writes a ``title`` (document heading, level 0) followed by
    ``sections`` in order to ``target_dir/filename_stem.docx``
    (collision-safe). Returns the final path."""
    import docx

    document = docx.Document()
    document.add_heading(title, level=0)

    for section in sections:
        if section.kind == "heading":
            document.add_heading(section.text, level=section.level)
        elif section.kind == "paragraph":
            document.add_paragraph(section.text)
        elif section.kind == "image":
            if section.image is None or not section.image.path:
                continue
            document.add_picture(section.image.path)
        elif section.kind == "table":
            if not section.columns:
                continue
            table = document.add_table(rows=1, cols=len(section.columns))
            for i, column_name in enumerate(section.columns):
                table.cell(0, i).text = str(column_name)
            for row in section.rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    if i < len(cells):
                        cells[i].text = str(value)

    target_dir.mkdir(parents=True, exist_ok=True)
    destination = unique_destination(target_dir / f"{filename_stem}.docx")
    document.save(str(destination))
    return destination


__all__ = ["DocxSection", "write_docx_document"]
